"""
pipeline/formatter.py — Word (.docx) Output Formatter

Exports humanized text to a Word document with proper formatting:
  - GOST 7.32-2017 (for academic/academic-essay register)
  - Free format (for journalistic/general/composition register)

Section detection: parses ## Heading → Heading 1, ### → Heading 2, #### → Heading 3
"""

from __future__ import annotations

from pathlib import Path
from typing import Optional

from loguru import logger


def export_to_docx(
    text: str,
    register: str,
    output_path: str | Path,
    config: dict,
    title: Optional[str] = None,
) -> Path:
    """
    Export text to a Word (.docx) file.

    Args:
        text: The humanized output text (may contain ## / ### markdown headings)
        register: Pipeline register ('academic', 'academic-essay', 'journalistic', 'general', ...)
        output_path: Destination .docx file path
        config: Full pipeline config dict (reads `formatter:` block)
        title: Optional document title (used as first Heading 1 if provided)

    Returns:
        Path to the created .docx file.
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt
    except ImportError:
        raise ImportError(
            "python-docx is required for Word export. Install with: pip install python-docx"
        )

    fmt = config.get("formatter", {})
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    is_gost = register in ("academic", "academic-essay")

    doc = Document()

    if is_gost:
        _apply_gost_page_setup(doc, fmt)
    else:
        _apply_free_page_setup(doc, fmt)

    # Set default paragraph style
    _configure_default_style(doc, fmt, is_gost)

    if title:
        _add_heading(doc, title, level=1, is_gost=is_gost, fmt=fmt)

    # Parse and write content
    _write_content(doc, text, is_gost=is_gost, fmt=fmt)

    doc.save(str(output_path))
    logger.info(f"Formatter: saved Word document to {output_path}")
    return output_path


def _apply_gost_page_setup(doc, fmt: dict) -> None:
    """Apply GOST 7.32-2017 page margins."""
    from docx.shared import Cm

    margins = fmt.get("gost_margins", {"left": 3.0, "right": 1.5, "top": 2.0, "bottom": 2.0})
    for section in doc.sections:
        section.left_margin = Cm(margins.get("left", 3.0))
        section.right_margin = Cm(margins.get("right", 1.5))
        section.top_margin = Cm(margins.get("top", 2.0))
        section.bottom_margin = Cm(margins.get("bottom", 2.0))


def _apply_free_page_setup(doc, fmt: dict) -> None:
    """Apply free-format page margins."""
    from docx.shared import Cm

    all_margin = fmt.get("free_margins", {}).get("all", 2.5)
    for section in doc.sections:
        section.left_margin = Cm(all_margin)
        section.right_margin = Cm(all_margin)
        section.top_margin = Cm(all_margin)
        section.bottom_margin = Cm(all_margin)


def _configure_default_style(doc, fmt: dict, is_gost: bool) -> None:
    """Configure the default paragraph style."""
    from docx.shared import Pt

    normal_style = doc.styles["Normal"]
    font = normal_style.font

    if is_gost:
        font.name = fmt.get("gost_font", "Times New Roman")
        font.size = Pt(fmt.get("gost_font_size", 14))
    else:
        font.name = fmt.get("free_font", "Calibri")
        font.size = Pt(fmt.get("free_font_size", 11))

    # Apply CJK/East Asian font name fallback
    _set_font_east_asian(normal_style, font.name)


def _set_font_east_asian(style, font_name: str) -> None:
    """Set East Asian font name in XML for proper rendering."""
    try:
        from docx.oxml.ns import qn
        from lxml import etree

        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:eastAsia"), font_name)
    except Exception:
        pass  # Non-critical; skip if lxml not available


def _add_heading(doc, text: str, level: int, is_gost: bool, fmt: dict) -> None:
    """Add a heading paragraph with correct style."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    style_name = f"Heading {level}"

    # Ensure the heading style exists
    try:
        heading = doc.add_heading(text, level=level)
    except Exception:
        heading = doc.add_paragraph(text)

    run = heading.runs[0] if heading.runs else heading.add_run(text)

    if is_gost:
        font_name = fmt.get("gost_font", "Times New Roman")
        font_size = fmt.get(f"gost_heading{level}_size", fmt.get("gost_font_size", 14))
        run.font.name = font_name
        run.font.size = Pt(font_size)

        if level == 1:
            run.font.bold = True
            # Chapter headings: uppercase
            if not text.isupper():
                run.text = text.upper()
        elif level == 2:
            run.font.bold = True
        else:
            run.font.bold = False

        # GOST: center chapter headings, left-align sub-headings
        if level == 1:
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        font_name = fmt.get("free_font", "Calibri")
        run.font.name = font_name
        if level == 1:
            run.font.bold = True
        elif level == 2:
            run.font.bold = True
        else:
            run.font.bold = False
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_body_paragraph(doc, text: str, is_gost: bool, fmt: dict) -> None:
    """Add a body paragraph with correct formatting."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
    from docx.oxml import OxmlElement

    para = doc.add_paragraph()
    run = para.add_run(text)

    if is_gost:
        font_name = fmt.get("gost_font", "Times New Roman")
        font_size = fmt.get("gost_font_size", 14)
        run.font.name = font_name
        run.font.size = Pt(font_size)

        # Line spacing: 1.5
        pf = para.paragraph_format
        pf.line_spacing = Pt(font_size * 1.5)  # Approximate 1.5 line spacing
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # First-line indent: 1.25 cm
        indent_cm = fmt.get("gost_indent_cm", 1.25)
        pf.first_line_indent = Cm(indent_cm)
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)

    else:
        font_name = fmt.get("free_font", "Calibri")
        font_size = fmt.get("free_font_size", 11)
        run.font.name = font_name
        run.font.size = Pt(font_size)

        pf = para.paragraph_format
        pf.line_spacing = fmt.get("free_line_spacing", 1.15)
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Cm(fmt.get("free_indent_cm", 0))
        pf.space_after = Pt(6)


def _write_content(doc, text: str, is_gost: bool, fmt: dict) -> None:
    """Parse text with markdown headings and write to document."""
    lines = text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]

        # Detect heading levels
        if line.startswith("#### "):
            _add_heading(doc, line[5:].strip(), level=3, is_gost=is_gost, fmt=fmt)
        elif line.startswith("### "):
            _add_heading(doc, line[4:].strip(), level=2, is_gost=is_gost, fmt=fmt)
        elif line.startswith("## "):
            _add_heading(doc, line[3:].strip(), level=1, is_gost=is_gost, fmt=fmt)
        elif line.startswith("# "):
            _add_heading(doc, line[2:].strip(), level=1, is_gost=is_gost, fmt=fmt)
        elif line.strip() == "":
            # Empty line: paragraph separator
            pass
        else:
            # Collect consecutive non-empty, non-heading lines as a paragraph
            para_lines = [line]
            while i + 1 < len(lines):
                next_line = lines[i + 1]
                if (
                    next_line.startswith("#")
                    or next_line.strip() == ""
                ):
                    break
                para_lines.append(next_line)
                i += 1
            para_text = " ".join(pl.strip() for pl in para_lines)
            if para_text.strip():
                _add_body_paragraph(doc, para_text, is_gost=is_gost, fmt=fmt)

        i += 1
